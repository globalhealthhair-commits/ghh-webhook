"""
Genera el PDF del informe a partir del template DOCX.
Variables en el template: {{NOMBRE}} y {{FOLICULOS}}
"""
import os, tempfile, subprocess
from docx import Document
import config

def generate_informe(nombre_completo: str, foliculos: str, tecnica: str = "FUE") -> str:
    """
    Rellena el template con los datos del paciente y devuelve la ruta al PDF generado.
    tecnica: 'FUE' o 'DHI'
    """
    doc = Document(config.TEMPLATE_PATH)

    # Sustituir variables en todos los runs
    subs = {
        "{{NOMBRE}}":    nombre_completo,
        "{{FOLICULOS}}": foliculos,
    }
    for para in doc.paragraphs:
        for run in para.runs:
            for old, new in subs.items():
                if old in run.text:
                    run.text = run.text.replace(old, new)

    # Si es DHI, destacar la mención DHI en el encabezado del informe
    if tecnica.upper() == "DHI":
        for para in doc.paragraphs:
            if "FUE ZAFIRO o D.H.I.(sin rasurar)" in para.text:
                for run in para.runs:
                    if "FUE ZAFIRO" in run.text:
                        run.text = run.text.replace(
                            "FUE ZAFIRO o D.H.I.(sin rasurar)",
                            "D.H.I. (sin rasurar) — técnica recomendada por el doctor"
                        )
                break

    # Guardar DOCX temporal
    tmp_dir = tempfile.mkdtemp()
    tmp_docx = os.path.join(tmp_dir, f"informe_{nombre_completo.replace(' ', '_')}.docx")
    doc.save(tmp_docx)

    # Convertir a PDF — intentar docx2pdf (usa Word en macOS), si falla usar LibreOffice
    pdf_path = tmp_docx.replace(".docx", ".pdf")
    try:
        from docx2pdf import convert
        convert(tmp_docx, pdf_path)
    except Exception:
        # Fallback: LibreOffice headless
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmp_dir, tmp_docx],
            capture_output=True, timeout=60
        )
        if result.returncode != 0:
            raise RuntimeError(f"PDF conversion failed: {result.stderr.decode()}")

    if not os.path.exists(pdf_path):
        raise FileNotFoundError(f"PDF no generado en {pdf_path}")

    return pdf_path
